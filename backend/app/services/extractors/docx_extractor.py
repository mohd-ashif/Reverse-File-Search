from pathlib import Path

from docx import Document

from app.services.extractors.base import TextExtractor


class DocxExtractor(TextExtractor):
    def extract(self, path: Path) -> str:
        document = Document(str(path))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs).strip()
